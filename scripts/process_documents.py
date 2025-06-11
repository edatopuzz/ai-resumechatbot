import os
from pathlib import Path
from openai import OpenAI
from dotenv import load_dotenv
from convex import ConvexClient
import docx
import json
from typing import List, Dict, Any
import time
from datetime import datetime
from tqdm import tqdm
import re

# Get the absolute path to the script's directory
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent

# Load environment variables
env_path = PROJECT_ROOT / ".env"
if not env_path.exists():
    raise FileNotFoundError(f"Environment file not found at {env_path}")
load_dotenv(env_path)

# Check for required environment variables
required_env_vars = ["OPENAI_API_KEY", "NEXT_PUBLIC_CONVEX_URL"]
missing_vars = [var for var in required_env_vars if not os.getenv(var)]
if missing_vars:
    raise ValueError(f"Missing required environment variables: {', '.join(missing_vars)}")

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Initialize Convex client
convex = ConvexClient(os.getenv("NEXT_PUBLIC_CONVEX_URL"))

def read_docx(file_path: str) -> Dict[str, Any]:
    """Read a .docx file and return its content with metadata."""
    try:
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract metadata
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "last_modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            "total_paragraphs": len(doc.paragraphs),
            "total_characters": len(content)
        }
        
        return {
            "content": content,
            "metadata": metadata
        }
    except Exception as e:
        print(f"Error reading {file_path}: {str(e)}")
        raise

def get_embedding(text: str) -> List[float]:
    """Get embedding for a text chunk using OpenAI API."""
    try:
        response = client.embeddings.create(
            input=text,
            model="text-embedding-3-small"  # Updated to newer model
        )
        return response.data[0].embedding
    except Exception as e:
        print(f"Error getting embedding: {str(e)}")
        raise

def calculate_semantic_similarity(text1: str, text2: str) -> float:
    """Calculate semantic similarity between two text segments using embeddings."""
    try:
        # Get embeddings for both texts
        embedding1 = get_embedding(text1)
        embedding2 = get_embedding(text2)
        
        # Calculate cosine similarity
        dot_product = sum(a * b for a, b in zip(embedding1, embedding2))
        magnitude1 = sum(a * a for a in embedding1) ** 0.5
        magnitude2 = sum(b * b for b in embedding2) ** 0.5
        
        if magnitude1 == 0 or magnitude2 == 0:
            return 0
        
        similarity = dot_product / (magnitude1 * magnitude2)
        return similarity
    except Exception as e:
        print(f"Error calculating similarity: {str(e)}")
        return 0

def split_into_semantic_chunks(text: str, similarity_threshold: float = 0.7, min_chunk_size: int = 200) -> List[Dict[str, Any]]:
    """Split text into semantic chunks based on content similarity."""
    try:
        # Split text into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if not sentences:
            return []
        
        chunks = []
        current_chunk = [sentences[0]]
        
        print(f"Processing {len(sentences)} sentences for semantic chunking...")
        
        for i in range(1, len(sentences)):
            current_sentence = sentences[i]
            
            # Get the last few sentences from current chunk for comparison
            comparison_text = ' '.join(current_chunk[-3:]) if len(current_chunk) >= 3 else ' '.join(current_chunk)
            
            # Calculate similarity between current sentence and chunk context
            try:
                similarity = calculate_semantic_similarity(comparison_text, current_sentence)
                print(f"Sentence {i+1}/{len(sentences)}: Similarity = {similarity:.3f}")
                
                # If similarity is high enough and chunk isn't too long, add to current chunk
                current_chunk_text = ' '.join(current_chunk)
                if similarity >= similarity_threshold and len(current_chunk_text) < 2000:
                    current_chunk.append(current_sentence)
                else:
                    # Start new chunk if current chunk meets minimum size
                    if len(' '.join(current_chunk)) >= min_chunk_size:
                        chunk_text = ' '.join(current_chunk)
                        chunks.append({
                            "content": chunk_text,
                            "metadata": {
                                "chunk_index": len(chunks),
                                "chunk_size": len(chunk_text),
                                "sentence_count": len(current_chunk),
                                "avg_similarity": similarity
                            }
                        })
                        current_chunk = [current_sentence]
                    else:
                        # If chunk is too small, add sentence anyway
                        current_chunk.append(current_sentence)
                        
            except Exception as e:
                print(f"Error processing sentence {i}: {str(e)}")
                # Fallback: just add to current chunk
                current_chunk.append(current_sentence)
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "metadata": {
                    "chunk_index": len(chunks),
                    "chunk_size": len(chunk_text),
                    "sentence_count": len(current_chunk)
                }
            })
        
        print(f"Created {len(chunks)} semantic chunks")
        return chunks
        
    except Exception as e:
        print(f"Error in semantic chunking: {str(e)}")
        # Fallback to simple chunking
        return simple_chunk_fallback(text)

def simple_chunk_fallback(text: str, chunk_size: int = 1000) -> List[Dict[str, Any]]:
    """Fallback to simple chunking if semantic chunking fails."""
    print("Falling back to simple chunking...")
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
            
        if current_size + len(paragraph) > chunk_size and current_chunk:
            chunk_text = '\n'.join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "metadata": {
                    "chunk_index": len(chunks),
                    "chunk_size": len(chunk_text),
                    "chunking_method": "fallback"
                }
            })
            current_chunk = []
            current_size = 0
        
        current_chunk.append(paragraph)
        current_size += len(paragraph)
    
    if current_chunk:
        chunk_text = '\n'.join(current_chunk)
        chunks.append({
            "content": chunk_text,
            "metadata": {
                "chunk_index": len(chunks),
                "chunk_size": len(chunk_text),
                "chunking_method": "fallback"
            }
        })
    
    return chunks

def process_documents():
    """Process all documents using semantic chunking and store them in the database."""
    try:
        # Get the absolute paths to the documents
        documents_dir = PROJECT_ROOT / "data"
        resume_path = documents_dir / "EdaTopuz_Resume 2025.docx"
        suggested_answers_path = documents_dir / "suggested_answers.docx"
        
        print(f"\n🧠 Starting document processing with SEMANTIC CHUNKING")
        print(f"📁 Looking for documents in {documents_dir}")
        
        # Process resume
        if resume_path.exists():
            print(f"\n📄 Processing resume: {resume_path.name}")
            resume_data = read_docx(str(resume_path))
            resume_content = resume_data["content"]
            
            if resume_content.strip():
                print(f"✅ Resume loaded: {len(resume_content)} characters")
                
                # Store full document
                resume_embedding = get_embedding(resume_content)
                resume_id = convex.mutation("documents:store", {
                    "name": "resume",
                    "content": resume_content,
                    "embedding": resume_embedding
                })
                print(f"📝 Stored full resume with ID: {resume_id}")
                
                # Create semantic chunks
                print("🧠 Creating semantic chunks...")
                semantic_chunks = split_into_semantic_chunks(resume_content)
                
                # Store semantic chunks
                for i, chunk in enumerate(semantic_chunks):
                    print(f"💾 Processing semantic chunk {i+1}/{len(semantic_chunks)}")
                    
                    embedding = get_embedding(chunk["content"])
                    result = convex.mutation("documents:store", {
                        "name": f"resume_semantic_chunk_{i}",
                        "content": chunk["content"],
                        "embedding": embedding
                    })
                    
                    print(f"✅ Stored semantic chunk {i+1} (size: {chunk['metadata']['chunk_size']} chars)")
                    time.sleep(0.5)  # Rate limiting
        
        # Process suggested answers
        if suggested_answers_path.exists():
            print(f"\n📄 Processing suggested answers: {suggested_answers_path.name}")
            answers_data = read_docx(str(suggested_answers_path))
            answers_content = answers_data["content"]
            
            if answers_content.strip():
                print(f"✅ Suggested answers loaded: {len(answers_content)} characters")
                
                # Store full document
                answers_embedding = get_embedding(answers_content)
                answers_id = convex.mutation("documents:store", {
                    "name": "suggested_answers",
                    "content": answers_content,
                    "embedding": answers_embedding
                })
                print(f"📝 Stored full suggested answers with ID: {answers_id}")
                
                # Create semantic chunks
                print("🧠 Creating semantic chunks...")
                semantic_chunks = split_into_semantic_chunks(answers_content)
                
                # Store semantic chunks
                for i, chunk in enumerate(semantic_chunks):
                    print(f"💾 Processing semantic chunk {i+1}/{len(semantic_chunks)}")
                    
                    embedding = get_embedding(chunk["content"])
                    result = convex.mutation("documents:store", {
                        "name": f"suggested_answers_semantic_chunk_{i}",
                        "content": chunk["content"],
                        "embedding": embedding
                    })
                    
                    print(f"✅ Stored semantic chunk {i+1} (size: {chunk['metadata']['chunk_size']} chars)")
                    time.sleep(0.5)  # Rate limiting
        
        print("\n🎉 SUCCESS! Documents processed with TRUE semantic chunking!")
        print("🧠 Chunks are now created based on semantic similarity, not arbitrary size!")
        print("📊 Each chunk groups semantically related content together")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        raise

if __name__ == "__main__":
    process_documents()
    